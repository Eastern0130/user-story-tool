import streamlit as st
import anthropic
import os
import base64
import io
from dotenv import load_dotenv
from PIL import Image
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

load_dotenv()

def _load_img(filename):
    path = os.path.join(os.path.dirname(__file__), filename)
    with open(path, "rb") as f:
        return f"data:image/png;base64,{base64.b64encode(f.read()).decode()}"

PARROT_SRC  = _load_img("parrot1.png")
PARROT_SRC2 = _load_img("parrot2.png")
PARROT_SRC3 = _load_img("parrot3.png")

_parrot_icon = Image.open(os.path.join(os.path.dirname(__file__), "parrot1.png"))
st.set_page_config(page_title="User Story 產生器", page_icon=_parrot_icon, layout="centered", initial_sidebar_state="collapsed")

st.markdown("""<style>
:root {
    --blue:#007AFF;--blue-hover:#0071E3;--blue-pressed:#0062CC;
    --blue-tint:rgba(0,122,255,0.10);--green:#34C759;--red:#FF3B30;--orange:#FF9500;
    --bg:#F2F2F7;--card:#FFFFFF;--label:#1C1C1E;--label-2:#6E6E73;--label-3:#AEAEB2;--sep:#E5E5EA;
    --r-md:12px;--r-lg:16px;
    --shadow:0 2px 8px rgba(0,0,0,0.07),0 0 1px rgba(0,0,0,0.05);
    --font:-apple-system,BlinkMacSystemFont,"SF Pro Text","SF Pro Display","Helvetica Neue",Arial,sans-serif;
}
html,body,[data-testid="stAppViewContainer"],[data-testid="stApp"],.stApp{background-color:var(--bg)!important;font-family:var(--font)!important;color:var(--label)!important;-webkit-font-smoothing:antialiased;}
[data-testid="stHeader"]{background-color:var(--bg)!important;border-bottom:1px solid var(--sep)!important;}
[data-testid="stMain"]{background-color:var(--bg)!important;}
[data-testid="stMainBlockContainer"],.block-container{padding-top:3.5rem!important;padding-bottom:4rem!important;max-width:700px!important;}
label,[data-testid="stWidgetLabel"] p{font-family:var(--font)!important;font-size:16px!important;font-weight:600!important;color:var(--label-2)!important;letter-spacing:0.03em!important;text-transform:uppercase!important;margin-bottom:6px!important;}
[data-testid="stTextArea"] textarea{font-family:var(--font)!important;font-size:18px!important;line-height:1.65!important;color:var(--label)!important;background-color:var(--card)!important;border:1.5px solid var(--sep)!important;border-radius:var(--r-md)!important;padding:14px 16px!important;box-shadow:none!important;transition:border-color 200ms ease,box-shadow 200ms ease!important;}
[data-testid="stTextArea"] textarea::placeholder{color:var(--label-3)!important;}
[data-testid="stTextArea"] textarea:focus,[data-testid="stTextArea"] div:focus-within,[data-baseweb="textarea"]:focus-within,[data-baseweb="base-input"]:focus-within{border-color:var(--blue)!important;box-shadow:0 0 0 3px var(--blue-tint)!important;outline:none!important;}
*:focus-visible{outline:2px solid var(--blue)!important;outline-offset:2px!important;}
[data-testid="stButton"]>button{font-family:var(--font)!important;font-size:18px!important;font-weight:800!important;color:#FFFFFF!important;background-color:var(--blue)!important;border:none!important;border-radius:980px!important;padding:13px 24px!important;letter-spacing:0.01em!important;box-shadow:0 1px 4px rgba(0,122,255,0.28)!important;transition:background-color 150ms ease,transform 100ms ease,box-shadow 150ms ease!important;}
[data-testid="stButton"]>button:hover{background-color:var(--blue-hover)!important;box-shadow:0 3px 12px rgba(0,122,255,0.32)!important;}
[data-testid="stButton"]>button:active{background-color:var(--blue-pressed)!important;transform:scale(0.975)!important;}
[data-testid="stAlert"]{border-radius:var(--r-md)!important;font-family:var(--font)!important;font-size:17px!important;}
[data-testid="stSuccess"]{background-color:rgba(52,199,89,0.10)!important;border-left:4px solid var(--green)!important;color:#1a5c2e!important;}
[data-testid="stWarning"]{background-color:rgba(255,149,0,0.10)!important;border-left:4px solid var(--orange)!important;color:#7a4800!important;}
[data-testid="stError"]{background-color:rgba(255,59,48,0.08)!important;border-left:4px solid var(--red)!important;color:#8b1a15!important;}
[data-testid="stSpinner"] p{color:var(--label-2)!important;font-family:var(--font)!important;font-size:17px!important;}
hr{border:none!important;border-top:1px solid var(--sep)!important;margin:1rem 0!important;}
.hero{display:flex;align-items:center;gap:20px;padding:0.5rem 0 0.6rem;}
.hero-parrot{height:120px;width:auto;object-fit:contain;flex-shrink:0;filter:drop-shadow(0 4px 12px rgba(0,0,0,0.15));}
.hero-text{flex:1;}
.hero-text h1{font-size:41px;font-weight:700;color:var(--label);margin:0 0 4px 0;letter-spacing:-0.03em;line-height:1.15;}
.hero-text p{font-size:20px;color:var(--label-2);margin:0;line-height:1.45;}
.section-label{display:flex;align-items:center;gap:8px;margin-bottom:10px;}
.step-dot{width:22px;height:22px;border-radius:50%;background:var(--blue);color:#fff;font-size:13px;font-weight:700;display:flex;align-items:center;justify-content:center;flex-shrink:0;}
.section-label-text{font-size:16px;font-weight:600;color:var(--label);}
.output-preview{display:flex;align-items:center;gap:8px;flex-wrap:wrap;margin:14px 0 20px;}
.preview-hint{font-size:14px;color:var(--label-3);font-weight:500;}
.preview-chip{display:inline-flex;align-items:center;gap:4px;font-size:14px;font-weight:600;padding:4px 10px;border-radius:980px;}
.chip-blue{background:#EBF3FF;color:#0055CC;}.chip-green{background:#E8F8ED;color:#1a6632;}.chip-orange{background:#FFF4E6;color:#7a4800;}
@keyframes fadeSlideUp{from{opacity:0;transform:translateY(10px);}to{opacity:1;transform:translateY(0);}}
.output-card{background:var(--card);border-radius:var(--r-lg);box-shadow:var(--shadow);padding:22px 26px;margin-bottom:16px;border:1px solid rgba(0,0,0,0.04);animation:fadeSlideUp 0.35s ease both;}
.output-card-header{margin-bottom:14px;padding-bottom:12px;border-bottom:1px solid var(--sep);}
.card-badge{display:inline-flex;align-items:center;gap:6px;font-size:14px;font-weight:700;padding:5px 12px;border-radius:980px;letter-spacing:0.02em;}
.badge-blue{background:#EBF3FF;color:#0055CC;}.badge-green{background:#E8F8ED;color:#1a6632;}.badge-orange{background:#FFF4E6;color:#7a4800;}
.output-card-body{font-family:var(--font);font-size:18px;color:var(--label);line-height:1.78;}
.output-card-body p{margin:0 0 10px 0;}
.output-card-body p:last-child{margin-bottom:0;}
.output-card-body strong,.output-card-body b{font-weight:600;color:var(--label);}
.loading-overlay{position:fixed;inset:0;background:rgba(242,242,247,0.36);backdrop-filter:blur(12px);-webkit-backdrop-filter:blur(12px);z-index:9999;display:flex;align-items:center;justify-content:center;}
.parrot-loading-wrap{background:linear-gradient(white,white) padding-box,linear-gradient(135deg,rgba(0,122,255,0.35),rgba(52,199,89,0.35)) border-box;border:2px solid transparent;border-radius:24px;box-shadow:0 12px 40px rgba(0,0,0,0.14),0 0 1px rgba(0,0,0,0.06),0 0 28px rgba(0,122,255,0.07);padding:2.64rem 3.6rem 2.16rem;text-align:center;display:flex;flex-direction:column;align-items:center;}
.parrot-loading{display:flex;justify-content:center;align-items:flex-end;gap:28px;margin-bottom:1.2rem;}
.loading-parrot{height:96px;width:auto;opacity:0;filter:saturate(1.3) brightness(1.05) drop-shadow(0 2px 6px rgba(0,0,0,0.10));}
.p1{animation:seq1 2.76s steps(1,end) infinite;}
.p2{animation:seq2 2.76s steps(1,end) infinite;}
.p3{animation:seq3 2.76s steps(1,end) infinite;}
@keyframes seq1{0%{opacity:1;}74.9%{opacity:1;}75%{opacity:0;}100%{opacity:0;}}
@keyframes seq2{0%{opacity:0;}24.9%{opacity:0;}25%{opacity:1;}74.9%{opacity:1;}75%{opacity:0;}100%{opacity:0;}}
@keyframes seq3{0%{opacity:0;}49.9%{opacity:0;}50%{opacity:1;}74.9%{opacity:1;}75%{opacity:0;}100%{opacity:0;}}
.loading-text{font-size:17px;color:var(--label-2);font-weight:500;letter-spacing:0.02em;}
.ac-item{display:flex;gap:14px;padding:14px 0;border-bottom:1px solid var(--sep);}
.ac-item:first-child{padding-top:0;}.ac-item:last-child{border-bottom:none;padding-bottom:0;}
.ac-num{font-size:20px;font-weight:700;color:var(--blue);min-width:22px;flex-shrink:0;line-height:1.5;}
.ac-body{flex:1;}
.ac-row{display:flex;align-items:baseline;gap:8px;margin-bottom:5px;}
.ac-row:last-child{margin-bottom:0;}
.ac-label{font-size:13px;font-weight:700;padding:2px 8px;border-radius:980px;letter-spacing:0.04em;flex-shrink:0;text-align:center;line-height:1.7;}
.label-premise{background:#F2F2F7;color:#6E6E73;}.label-action{background:#EBF3FF;color:#0055CC;}.label-result{background:#E8F8ED;color:#1a6632;}.label-risk{background:#FFF4E6;color:#7a4800;}.label-clarify{background:#F2F0FF;color:#4B35A1;}
.ac-text{font-size:17px;color:var(--label);line-height:1.65;}
.page-footer{text-align:center;font-size:14px;color:var(--label-3);line-height:1.7;padding:0.5rem 0;}
@media(max-width:640px){.hero h1{font-size:31px!important;}.output-card{padding:18px 20px!important;}[data-testid="stMainBlockContainer"],.block-container{padding-left:1rem!important;padding-right:1rem!important;}}
section[data-testid="stSidebar"]>div:first-child{padding-top:0.8rem!important;}
[data-testid="stSidebar"] h3{width:100%!important;display:block!important;font-size:20px!important;font-weight:800!important;padding:0.2rem 0 0.6rem!important;border-bottom:1px solid var(--sep)!important;margin:0 0 0.6rem 0!important;}
[data-testid="stSidebar"] [data-testid="stButton"]>button[kind="secondary"]{background:transparent!important;border:none!important;color:var(--label)!important;text-align:left!important;box-shadow:none!important;font-size:13px!important;font-weight:500!important;padding:3px 6px!important;border-radius:4px!important;}
[data-testid="stSidebar"] [data-testid="stButton"]>button[kind="secondary"]:hover{background:rgba(0,0,0,0.05)!important;color:var(--label)!important;box-shadow:none!important;}
[data-testid="stSidebar"] [data-testid="stButton"]>button[kind="primary"]{font-size:13px!important;font-weight:700!important;padding:5px 12px!important;border-radius:980px!important;}
[data-testid="stSidebar"] [data-baseweb="checkbox"] input:checked+div,[data-testid="stSidebar"] [role="checkbox"][aria-checked="true"]{background-color:#AEAEB2!important;border-color:#AEAEB2!important;}
[data-testid="stStatusWidget"]{display:none!important;}
</style>""", unsafe_allow_html=True)

if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

if not st.session_state.authenticated:
    st.markdown("""<style>
    [data-testid="stForm"]{background:#FFFFFF;border-radius:20px;box-shadow:0 2px 8px rgba(0,0,0,0.07),0 0 1px rgba(0,0,0,0.05);border:none!important;padding:3rem 3.5rem 2.8rem;}
    [data-testid="InputInstructions"]{display:none!important;}
    [data-testid="stForm"] [data-baseweb="input"]{border:1.5px solid #BDD7FF!important;border-radius:10px!important;background:#FFFFFF!important;}
    [data-testid="stForm"] [data-baseweb="input"]:focus-within{border-color:var(--blue)!important;box-shadow:0 0 0 3px var(--blue-tint)!important;}
    [data-testid="stForm"] [data-baseweb="input"] input{border:none!important;border-right:none!important;outline:none!important;box-shadow:none!important;}
    [data-testid="stForm"] [data-baseweb="input"]>div{border:none!important;border-left:none!important;background:transparent!important;}
    [data-testid="stForm"] [data-baseweb="input"] button{border:none!important;border-right:none!important;background:transparent!important;box-shadow:none!important;outline:none!important;}
    [data-testid="stForm"] [data-baseweb="input"] *{border-color:transparent!important;}
    [data-testid="stForm"] [data-testid="stFormSubmitButton"]>button{font-family:var(--font)!important;font-size:17px!important;font-weight:600!important;color:#FFFFFF!important;background-color:var(--blue)!important;border:none!important;border-radius:980px!important;padding:11px 24px!important;box-shadow:0 1px 4px rgba(0,122,255,0.28)!important;}
    </style>""", unsafe_allow_html=True)
    st.markdown("<div style='height:10vh'></div>", unsafe_allow_html=True)
    _, col, _ = st.columns([1, 4, 1])
    with col:
        with st.form("login_form"):
            st.markdown(f"""
            <div style="text-align:center;padding:0.5rem 0 1.4rem;">
                <img src="{PARROT_SRC}" style="height:110px;width:auto;filter:drop-shadow(0 4px 12px rgba(0,0,0,0.15));margin-bottom:1.2rem;">
                <div style="font-size:23px;font-weight:700;color:#1C1C1E;margin-bottom:0.4rem;">User Story 產生器</div>
                <div style="font-size:15px;color:#6E6E73;margin-bottom:0.4rem;">請輸入密碼以繼續</div>
            </div>""", unsafe_allow_html=True)
            pwd = st.text_input("", type="password", placeholder="密碼", label_visibility="collapsed")
            submitted = st.form_submit_button("進入", use_container_width=True)
        if submitted:
            if pwd == os.environ.get("APP_PASSWORD", ""):
                st.session_state.authenticated = True
                st.rerun()
            else:
                st.error("密碼錯誤")
    st.components.v1.html("""<script>
    setTimeout(function(){
        var inp=window.parent.document.querySelector('input[type="password"]');
        if(inp){inp.setAttribute('autocomplete','new-password');inp.focus();}
    },300);
    </script>""", height=0)
    st.stop()

if "history" not in st.session_state:
    st.session_state.history = []
if "results" not in st.session_state:
    st.session_state.results = None
if "input_key" not in st.session_state:
    st.session_state.input_key = 0

def _set_font(run, name="標楷體", size=12, bold=False):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:eastAsia", "w:ascii", "w:hAnsi"):
        rFonts.set(qn(attr), name)

def _cell_text(cell, text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    _set_font(p.add_run(text), size=size, bold=bold)

def _cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def _center_cell(cell):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def generate_word(entries):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)
    numbers = ["①","②","③","④","⑤","⑥"]
    for idx, entry in enumerate(entries):
        if idx > 0:
            doc.add_page_break()
        t = doc.add_table(rows=1, cols=2); t.style = "Table Grid"
        t.columns[0].width = Cm(2.0)
        _cell_text(t.rows[0].cells[0], "需求描述", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER); _cell_bg(t.rows[0].cells[0], "F2F2F7"); _center_cell(t.rows[0].cells[0])
        _cell_text(t.rows[0].cells[1], entry["requirement"])
        doc.add_paragraph()
        t2 = doc.add_table(rows=2, cols=1); t2.style = "Table Grid"
        _cell_text(t2.rows[0].cells[0], "User Story", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER); _cell_bg(t2.rows[0].cells[0], "EBF3FF"); _center_cell(t2.rows[0].cells[0])
        _cell_text(t2.rows[1].cells[0], entry["user_story"])
        doc.add_paragraph()
        ac_items = [i.strip() for i in entry["ac_block"].split("\n\n") if i.strip()]
        t3 = doc.add_table(rows=1+len(ac_items), cols=4); t3.style = "Table Grid"
        t3.columns[0].width = Cm(0.9)
        for cell, lbl in zip(t3.rows[0].cells, ["#","前提","操作","預期結果"]):
            _cell_text(cell, lbl, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER); _cell_bg(cell, "E8F8ED"); _center_cell(cell)
        for i, item in enumerate(ac_items):
            d = {"前提":"","操作":"","預期":""}
            for line in [l.strip() for l in item.split("\n") if l.strip()]:
                for k in d:
                    if line.startswith(k): d[k] = line.split("：",1)[-1].strip()
            r = t3.rows[i+1].cells
            _cell_text(r[0], numbers[i] if i<len(numbers) else f"{i+1}.", align=WD_ALIGN_PARAGRAPH.CENTER); _center_cell(r[0])
            _cell_text(r[1], d["前提"]); _cell_text(r[2], d["操作"]); _cell_text(r[3], d["預期"])
        doc.add_paragraph()
        risk_items = [i.strip() for i in entry["risk_block"].split("\n\n") if i.strip()]
        t4 = doc.add_table(rows=1+len(risk_items), cols=3); t4.style = "Table Grid"
        t4.columns[0].width = Cm(0.9)
        for cell, lbl in zip(t4.rows[0].cells, ["#","風險描述","釐清問題"]):
            _cell_text(cell, lbl, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER); _cell_bg(cell, "FFF4E6"); _center_cell(cell)
        for i, item in enumerate(risk_items):
            d = {"描述":"","釐清":""}
            for line in [l.strip() for l in item.split("\n") if l.strip()]:
                for k in d:
                    if line.startswith(k): d[k] = line.split("：",1)[-1].strip()
            r = t4.rows[i+1].cells
            _cell_text(r[0], numbers[i] if i<len(numbers) else f"{i+1}.", align=WD_ALIGN_PARAGRAPH.CENTER); _center_cell(r[0])
            _cell_text(r[1], d["描述"]); _cell_text(r[2], d["釐清"])
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.getvalue()

with st.sidebar:
    st.markdown('<div style="background:#EBF3FF;border-radius:10px;padding:0.55rem 0.9rem;margin-bottom:0.8rem;"><span style="font-size:17px;font-weight:800;color:#0055CC;">歷史紀錄</span></div>', unsafe_allow_html=True)
    if not st.session_state.history:
        st.caption("產生後將顯示於此")
    else:
        n = len(st.session_state.history)
        selected = []
        for real_idx in range(n-1, -1, -1):
            entry = st.session_state.history[real_idx]
            preview = entry["requirement"][:35] + "…" if len(entry["requirement"]) > 35 else entry["requirement"]
            cb_col, txt_col = st.columns([1, 6])
            with cb_col:
                checked = st.checkbox("", key=f"hist_{real_idx}")
            with txt_col:
                if st.button(preview, key=f"view_{real_idx}", use_container_width=True):
                    st.session_state.results = {
                        "user_story": entry["user_story"],
                        "ac_block": entry["ac_block"],
                        "risk_block": entry["risk_block"],
                    }
                    st.rerun()
            if checked:
                selected.append(real_idx)
        st.divider()
        c1, c2 = st.columns(2)
        if c1.button("全選", use_container_width=True, type="primary"):
            for i in range(n): st.session_state[f"hist_{i}"] = True
        if c2.button("清除", use_container_width=True, type="primary"):
            for i in range(n): st.session_state[f"hist_{i}"] = False
        if selected:
            sel_entries = [st.session_state.history[i] for i in sorted(selected)]
            doc_bytes = generate_word(sel_entries)
            st.download_button(
                label=f"⬇ 匯出 Word（{len(selected)} 份）",
                data=doc_bytes,
                file_name=f"user_story_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )

st.markdown(f"""
<div class="hero">
    <img src="{PARROT_SRC}" class="hero-parrot" alt="鸚鵡">
    <div class="hero-text">
        <h1>User Story 產生器</h1>
        <p>比起自己通靈，不如讓我來幫你</p>
    </div>
</div>""", unsafe_allow_html=True)

st.markdown("<hr>", unsafe_allow_html=True)

st.markdown("""
<div class="section-label">
    <div class="step-dot">1</div>
    <span class="section-label-text">輸入功能需求描述</span>
</div>""", unsafe_allow_html=True)

requirement = st.text_area(
    label="功能需求描述",
    placeholder="例如：使用者可以在登入後修改自己的個人資料，包含姓名、電話與 Email，修改後需要寄送確認信",
    height=160,
    label_visibility="collapsed",
    key=f"req_{st.session_state.input_key}",
)

st.markdown("""
<div class="output-preview">
    <span class="preview-hint">將產出：</span>
    <span class="preview-chip chip-blue">📖 User Story</span>
    <span class="preview-chip chip-green">✅ Acceptance Criteria</span>
    <span class="preview-chip chip-orange">⚠️ Risk Analysis</span>
</div>""", unsafe_allow_html=True)

def extract_block(text, start_tag, end_tag):
    try:
        return text.split(start_tag)[1].split(end_tag)[0].strip()
    except IndexError:
        return ""

def to_html(text):
    return "".join(f"<p>{p.replace(chr(10),'<br>')}</p>" for p in text.split("\n\n"))

def risk_to_html(text):
    items = [i.strip() for i in text.split("\n\n") if i.strip()]
    numbers = ["①","②","③","④","⑤","⑥"]
    html_parts = []
    idx = 0
    for item in items:
        rows = []
        for line in [l.strip() for l in item.split("\n") if l.strip()]:
            if line.startswith("風險"):
                continue
            if line.startswith("描述"):
                rows.append(f'<div class="ac-row"><span class="ac-label label-risk">風險</span><span class="ac-text">{line.split("：",1)[-1].strip()}</span></div>')
            elif line.startswith("釐清"):
                rows.append(f'<div class="ac-row"><span class="ac-label label-clarify">釐清</span><span class="ac-text">{line.split("：",1)[-1].strip()}</span></div>')
            else:
                rows.append(f'<div class="ac-row"><span class="ac-text">{line}</span></div>')
        if rows:
            num = numbers[idx] if idx < len(numbers) else f"{idx+1}."
            html_parts.append(f'<div class="ac-item"><div class="ac-num">{num}</div><div class="ac-body">{"".join(rows)}</div></div>')
            idx += 1
    return "".join(html_parts)

def ac_to_html(text):
    items = [i.strip() for i in text.split("\n\n") if i.strip()]
    numbers = ["①","②","③","④","⑤","⑥"]
    html_parts = []
    idx = 0
    for item in items:
        rows = []
        for line in [l.strip() for l in item.split("\n") if l.strip()]:
            if line.startswith("條件"):
                continue
            if line.startswith("前提"):
                rows.append(f'<div class="ac-row"><span class="ac-label label-premise">前提</span><span class="ac-text">{line.split("：",1)[-1].strip()}</span></div>')
            elif line.startswith("操作"):
                rows.append(f'<div class="ac-row"><span class="ac-label label-action">操作</span><span class="ac-text">{line.split("：",1)[-1].strip()}</span></div>')
            elif line.startswith("預期"):
                rows.append(f'<div class="ac-row"><span class="ac-label label-result">預期</span><span class="ac-text">{line.split("：",1)[-1].strip()}</span></div>')
            else:
                rows.append(f'<div class="ac-row"><span class="ac-text">{line}</span></div>')
        if rows:
            num = numbers[idx] if idx < len(numbers) else f"{idx+1}."
            html_parts.append(f'<div class="ac-item"><div class="ac-num">{num}</div><div class="ac-body">{"".join(rows)}</div></div>')
            idx += 1
    return "".join(html_parts)

if st.button("產生 User Story", type="primary", use_container_width=True):
    if not requirement.strip():
        st.warning("請先輸入需求描述，再按產生按鈕")
    else:
        loading_ph = st.empty()
        loading_ph.markdown(f"""
        <div class="loading-overlay">
            <div class="parrot-loading-wrap">
                <div class="parrot-loading">
                    <img src="{PARROT_SRC2}" class="loading-parrot p1" alt="">
                    <img src="{PARROT_SRC2}" class="loading-parrot p2" alt="">
                    <img src="{PARROT_SRC2}" class="loading-parrot p3" alt="">
                </div>
                <div class="loading-text">載入中...</div>
            </div>
        </div>""", unsafe_allow_html=True)

        try:
            client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))
            prompt = f"""你是一位資深的軟體專案需求分析師。請根據以下功能需求描述，產出結構化的需求文件。

功能需求描述：
{requirement}

請輸出以下三個部分，每個部分用指定的標記包起來，不要加其他標題或說明：

[USER_STORY_START]
（使用「身為＿＿，我希望能夠＿＿，以便＿＿」格式，一段完整句子）
[USER_STORY_END]

[AC_START]
每條驗收條件請使用以下格式，條件之間空一行，不要加其他符號或說明：

條件 1
前提：xxx
操作：xxx
預期結果：xxx

（依此類推，共 3~6 條）
[AC_END]

[RISK_START]
每條風險請使用以下格式，風險之間空一行，不要加其他符號或說明：

風險 1
描述：xxx
釐清：xxx

（依此類推，共 2~4 條）
[RISK_END]

請用繁體中文回答，語氣專業清晰。"""

            message = client.messages.create(
                model="claude-sonnet-4-6",
                max_tokens=1800,
                messages=[{"role": "user", "content": prompt}]
            )
            result     = message.content[0].text
            user_story = extract_block(result, "[USER_STORY_START]", "[USER_STORY_END]")
            ac_block   = extract_block(result, "[AC_START]",         "[AC_END]")
            risk_block = extract_block(result, "[RISK_START]",       "[RISK_END]")

            loading_ph.empty()
            st.session_state.history.append({
                "time": datetime.now().strftime("%H:%M"),
                "requirement": requirement,
                "user_story": user_story,
                "ac_block": ac_block,
                "risk_block": risk_block,
            })
            st.session_state.results = {"user_story": user_story, "ac_block": ac_block, "risk_block": risk_block}
            st.rerun()

        except anthropic.AuthenticationError:
            loading_ph.empty()
            st.error("API 金鑰錯誤，請確認 .env 檔案中的 ANTHROPIC_API_KEY 是否填寫正確")
        except anthropic.RateLimitError:
            loading_ph.empty()
            st.error("API 使用量已達上限，請稍後再試")
        except Exception as e:
            loading_ph.empty()
            st.error(f"發生未預期的錯誤：{str(e)}")

if st.session_state.results:
    r = st.session_state.results
    st.success("分析完成")
    st.markdown("<div style='height:4px'></div>", unsafe_allow_html=True)
    st.markdown(f'<div class="output-card" style="animation-delay:0s"><div class="output-card-header"><span class="card-badge badge-blue">📖 User Story</span></div><div class="output-card-body">{to_html(r["user_story"])}</div></div>', unsafe_allow_html=True)
    st.markdown(f'<div class="output-card" style="animation-delay:0.1s"><div class="output-card-header"><span class="card-badge badge-green">✅ Acceptance Criteria</span></div><div class="output-card-body">{ac_to_html(r["ac_block"])}</div></div>', unsafe_allow_html=True)
    st.markdown(f'<div class="output-card" style="animation-delay:0.2s"><div class="output-card-header"><span class="card-badge badge-orange">⚠️ Risk Analysis</span></div><div class="output-card-body">{risk_to_html(r["risk_block"])}</div></div>', unsafe_allow_html=True)
    st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)
    if st.button("下一則", use_container_width=True):
        st.session_state.results = None
        st.session_state.input_key += 1
        st.rerun()

st.markdown("<hr>", unsafe_allow_html=True)
st.markdown('<div class="page-footer">本工具在本機運行 · 輸入內容僅傳送至 Anthropic API 進行分析<br>不會儲存於任何第三方系統</div>', unsafe_allow_html=True)
